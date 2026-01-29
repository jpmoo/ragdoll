"""Extract plain text or structured document (text/chart/table blocks) from supported file types."""

import logging
from dataclasses import dataclass, field
from pathlib import Path

from . import config

logger = logging.getLogger(__name__)


# --- Structured extraction for PDF, DOCX, Excel (RAG guide: separate prose, charts, tables) ---

@dataclass
class TextBlock:
    page: int | None
    text: str


@dataclass
class ChartRegion:
    page: int
    image_bytes: bytes
    image_ext: str  # e.g. "png", "jpeg"


@dataclass
class TableRegion:
    page: int | None
    data: list[list[str]]  # rows of cells


@dataclass
class FigureRegion:
    page: int
    image_bytes: bytes  # rendered page or crop


@dataclass
class ImageRegion:
    """Embedded image (e.g. DOCX) to be classified and routed."""
    image_bytes: bytes
    ext: str
    page_or_idx: int | None  # page for PDF, index for DOCX


@dataclass
class Document:
    text_blocks: list[TextBlock] = field(default_factory=list)
    chart_regions: list[ChartRegion] = field(default_factory=list)
    table_regions: list[TableRegion] = field(default_factory=list)
    figure_regions: list[FigureRegion] = field(default_factory=list)
    image_regions: list[ImageRegion] = field(default_factory=list)

    def has_embeddable(self) -> bool:
        return bool(
            self.text_blocks or self.chart_regions or self.table_regions
            or self.figure_regions or self.image_regions
        )


def extract_document(path: Path) -> Document | None:
    """
    Structured extraction for PDF, DOCX, Excel. Returns None for plain text, images, or unsupported.
    When RAGDOLL_USE_DOCLING=true and docling is installed, tries Docling first for PDF/DOCX/XLSX/PPTX/image; falls back to legacy on failure.
    - PDF: text blocks per page; low-text pages with images -> chart regions.
    - DOCX: paragraphs -> text; doc.tables -> table regions.
    - Excel: each sheet -> one table region.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    # Optional: try Docling first when enabled (pip install '.[docling]', RAGDOLL_USE_DOCLING=true)
    if config.USE_DOCLING and suffix in _docling_ext():
        try:
            from .extractors_docling import extract_document_with_docling
            doc = extract_document_with_docling(path)
            if doc is not None and doc.has_embeddable():
                return doc
        except Exception as e:
            logger.debug("Docling extraction failed, using legacy: %s", e)

    if suffix in config.PDF_EXT:
        return _extract_pdf_document(path)
    if suffix in config.WORD_EXT:
        return _extract_docx_document(path)
    if suffix in config.EXCEL_EXT:
        return _extract_excel_document(path)
    return None


def _docling_ext() -> set:
    """Extensions Docling can ingest (avoids importing extractors_docling at module load)."""
    return config.PDF_EXT | config.WORD_EXT | config.EXCEL_EXT | {".pptx"} | config.IMAGE_EXT


def _page_to_png_bytes(page) -> bytes:
    """Render a fitz page to PNG bytes."""
    import io
    pix = page.get_pixmap(dpi=150)
    if hasattr(pix, "tobytes") and callable(getattr(pix, "tobytes", None)):
        try:
            return pix.tobytes(output="png")
        except Exception:
            pass
    from PIL import Image
    try:
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    except Exception:
        img = Image.frombytes("RGBA", [pix.width, pix.height], pix.samples).convert("RGB")
    buf = io.BytesIO()
    img.save(buf, "PNG")
    return buf.getvalue()


def _extract_pdf_document(path: Path) -> Document:
    import fitz

    doc = fitz.open(path)
    out = Document()
    try:
        # PDF table detection via pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(path) as pb:
                for i, pg in enumerate(pb.pages):
                    for t in (pg.find_tables() or []):
                        rows = t.extract()
                        if not rows:
                            continue
                        nonempty = sum(1 for r in rows for c in r if (c or "").strip())
                        # Skip tables with ≤2 non-empty cells (layout/grid noise on diagrams)
                        if nonempty <= 2:
                            continue
                        if any(any(c for c in r) for r in rows):
                            out.table_regions.append(TableRegion(page=i + 1, data=rows))
        except Exception as e:
            logger.debug("pdfplumber table extraction: %s", e)

        for pagenum in range(len(doc)):
            page = doc[pagenum]
            page_num = pagenum + 1
            raw = page.get_text()
            images = page.get_images(full=True)
            # Many short blocks or vector drawings -> figure (boxes, arrows, labels)
            blocks = (page.get_text("dict") or {}).get("blocks") or []
            short = [b for b in blocks if sum(len(s.get("text", "")) for l in b.get("lines", []) for s in l.get("spans", [])) < 80]
            try:
                drawings = page.get_drawings() or []
            except Exception:
                drawings = []

            # Prefer figure when we see diagram signals. Check before chart so figures
            # with small embedded images still get process interpretation.
            # - drawings + <1000 chars: strong diagram signal (arrows, boxes)
            # - no drawings: <800 chars and >=1 short block (labels in small blocks)
            is_figure = (drawings and len(raw.strip()) < 1000) or (
                len(raw.strip()) < 800 and len(short) >= 1
            )
            if is_figure:
                try:
                    out.figure_regions.append(
                        FigureRegion(page=page_num, image_bytes=_page_to_png_bytes(page))
                    )
                except Exception as e:
                    logger.warning("Could not render figure page %s: %s", page_num, e)
            # Low text + has images (and not figure above) -> chart regions
            elif len(raw.strip()) < 200 and images:
                for im in images:
                    xref = im[0]
                    try:
                        img = doc.extract_image(xref)
                        out.chart_regions.append(
                            ChartRegion(page=page_num, image_bytes=img["image"], image_ext=img["ext"] or "png")
                        )
                    except Exception as e:
                        logger.warning("Could not extract image xref=%s on page %s: %s", xref, page_num, e)
            elif raw.strip():
                out.text_blocks.append(TextBlock(page=page_num, text=raw))

        # Drop pdfplumber tables on figure pages; they are usually layout noise, not real tables.
        figure_pages = {f.page for f in out.figure_regions}
        out.table_regions = [t for t in out.table_regions if t.page not in figure_pages]
    finally:
        doc.close()
    return out


def _extract_docx_document(path: Path) -> Document:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    out = Document()
    for p in doc.paragraphs:
        if p.text.strip():
            out.text_blocks.append(TextBlock(page=None, text=p.text))
    for t in doc.tables:
        data = [[(c.text or "").strip() for c in row.cells] for row in t.rows]
        if any(any(c for c in row) for row in data):
            out.table_regions.append(TableRegion(page=None, data=data))

    # Embedded images: rels -> image parts
    try:
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" not in str(getattr(rel, "reltype", "")):
                continue
            part = getattr(rel, "target_part", None)
            if part is None or not hasattr(part, "blob"):
                continue
            blob = part.blob
            ct = getattr(part, "content_type", "") or ""
            ext = "png"
            if "jpeg" in ct or "jpg" in ct:
                ext = "jpeg"
            elif "gif" in ct:
                ext = "gif"
            elif "bmp" in ct:
                ext = "bmp"
            out.image_regions.append(ImageRegion(image_bytes=blob, ext=ext, page_or_idx=i))
    except Exception as e:
        logger.debug("DOCX embedded images: %s", e)
    return out


def _extract_excel_document(path: Path) -> Document:
    out = Document()
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        import openpyxl

        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        for sh in wb.worksheets:
            data = [
                [str(c.value).strip() if c.value is not None else "" for c in row]
                for row in sh.iter_rows(max_row=500)
            ]
            data = [r for r in data if any(r)]  # drop empty rows
            if data:
                out.table_regions.append(TableRegion(page=None, data=data))
        if hasattr(wb, "close"):
            wb.close()
    else:
        import xlrd

        wb = xlrd.open_workbook(path)
        for i in range(wb.nsheets):
            sh = wb.sheet_by_index(i)
            data = [[str(sh.cell_value(r, c)) for c in range(sh.ncols)] for r in range(min(sh.nrows, 500))]
            data = [r for r in data if any(r)]
            if data:
                out.table_regions.append(TableRegion(page=None, data=data))
    return out


def ocr_image_bytes(image_bytes: bytes) -> str:
    """OCR on raw image bytes (e.g. chart region). Returns extracted text."""
    import io
    import pytesseract
    from PIL import Image

    img = Image.open(io.BytesIO(image_bytes))
    return pytesseract.image_to_string(img) or ""


# --- Plain text extraction (fallback and for .txt, .md, .image) ---

def extract_text(path: Path) -> str:
    """Extract text from a file. Raises ValueError if unsupported or on error."""
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in config.TEXT_EXT:
        return _extract_plain(path)
    if suffix in config.WORD_EXT:
        return _extract_docx(path)
    if suffix in config.EXCEL_EXT:
        return _extract_excel(path)
    if suffix in config.PDF_EXT:
        return _extract_pdf(path)
    if suffix in config.IMAGE_EXT:
        return _extract_image(path)

    raise ValueError(f"Unsupported extension: {suffix}")


def _extract_plain(path: Path) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def _extract_excel(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        import openpyxl

        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    else:
        import xlrd

        wb = xlrd.open_workbook(path)

    parts: list[str] = []
    if hasattr(wb, "sheetnames"):
        for name in wb.sheetnames:
            sh = wb[name]
            for row in sh.iter_rows(values_only=True):
                line = "\t".join(str(c) if c is not None else "" for c in row).strip()
                if line:
                    parts.append(line)
    else:
        for i in range(wb.nsheets):
            sh = wb.sheet_by_index(i)
            for r in range(sh.nrows):
                row = [sh.cell_value(r, c) for c in range(sh.ncols)]
                line = "\t".join(str(c) if c else "" for c in row).strip()
                if line:
                    parts.append(line)

    if hasattr(wb, "close"):
        wb.close()
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def _extract_image(path: Path) -> str:
    import pytesseract
    from PIL import Image

    img = Image.open(path)
    return pytesseract.image_to_string(img)
